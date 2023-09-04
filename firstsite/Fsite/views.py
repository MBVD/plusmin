from codecs import ascii_encode
from datetime import datetime
from multiprocessing import context
from re import template
import re
from docx.shared import RGBColor
from urllib import response
from django.contrib.auth import logout, login
from django.contrib.auth.decorators import login_required
from django.contrib.auth.views import LoginView
from django.core.paginator import Paginator
from django.http import HttpResponse, HttpResponseNotFound, Http404
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse_lazy
from django.views.generic import ListView, DetailView, CreateView
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth import get_user
from django.contrib.auth.models import User
from django.core import serializers
import json
from .forms import *
from .models import *
from .utils import *
import random
import os
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.db.models import Q
from .filters import *
from django.db.models import Sum
from django.core.exceptions import ObjectDoesNotExist
from django.db.models.query import QuerySet
from django.forms import formset_factory
from django.forms import modelformset_factory
from io import BytesIO
from django.http import FileResponse
from docx import Document
from django.conf import settings
from docx.enum.text import WD_ALIGN_PARAGRAPH
# !/usr/bin/env python3


def random_questions(n, query, model):
    maked_query = model.objects.none()
    default_len = len(query)
    while default_len - len(query) < n:
        a = random.randint(0, len(query)-1)
        random_question = query[a]
        maked_query = maked_query | model.objects.filter(pk=random_question.pk)
        query = query.exclude(pk=random_question.pk)
    return maked_query.all()


def transform(argument):
    return datetime.strftime(argument, "%Y-%m-%d %H:%M")


def index(request):
    friends_count = 0
    interesting_news = News.objects.filter(is_published = True, tags__slug__in=['interesting_news'])
    random_news = None
    if interesting_news:
        random_news = interesting_news[random.randint(0, len(interesting_news)-1)]
    if (request.user.is_authenticated):
        friends_count = UserFollowing.objects.filter(user=request.user).count() 
    rating = UserProfile.objects.all().order_by('-rate')[:10]
    return render(request, "Fsite/index.html", {'title': "Главная",
                                                'alerts': News.objects.all().order_by('-time_update'), 
                                                'friends_count': friends_count, 
                                                'random_news': random_news,
                                                'rating': rating})


def problems(request):
    posts = Problem.objects.all().order_by('-rate')
    my_filter = ProblemsFilters(request.GET, queryset=posts)
    posts = my_filter.qs
    paginator = Paginator(posts, 20)
    page_number = request.GET.get('page', 1)
    subjects = MaterialSubject.objects.exclude(name='history')
    materials = Material.objects.none()
    for subject in subjects:
        if request.GET.get(subject.name, None):
            materials = materials | Material.objects.filter(subject__name=subject.name, content=request.GET.get(subject.name, None))
    history_year = request.GET.get('history_year', None)
    history_content = request.GET.get('history', None)
    history_questions_num = request.GET.get('history_questions_number', None)
    history_test_questions_num = request.GET.get('history_test_questions_number', None)
    # надо вытащить рандомные задачи
    history_questions_all = HistoryProblem.objects.filter(is_published=True).all()
    history_questions = HistoryProblem.objects.none()
    history_test_questions = HistoryTestProblem.objects.none()
    history_test_questions_all = HistoryTestProblem.objects.filter(is_published=True).all()
    if history_questions_num:
        history_questions_num = int(history_questions_num)
        if history_questions_num < len(history_questions_all):
            history_questions = random_questions(history_questions_num, history_questions_all, HistoryProblem)
        else:
            history_questions = history_questions_all.all()
    if history_test_questions_num:
        history_test_questions_num = int(history_test_questions_num)
        if history_test_questions_num < len(history_test_questions_all):
            history_test_questions = random_questions(history_test_questions_num, history_test_questions_all, HistoryTestProblem)
        else:
            history_test_questions = history_test_questions_all.all()
    history = Material.objects.none()
    if history_year:
        history = Material.objects.filter(subject__name='history', time_start__lte=history_year, time_end__gte=history_year)
    if history_content:
        if history_year:
            history = history & Material.objects.filter(subject__name='history', content__iregex=history_content)
        else:
            history = Material.objects.filter(subject__name='history', content__iregex=history_content)
    try:
        page_obj = paginator.get_page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.get_page(1)
    except EmptyPage:
        page_obj = paginator.get_page(paginator.num_pages)
    friends_count = 0
    interesting_news = News.objects.filter(is_published = True, tags__slug__in=['interesting_news'])
    random_news = None
    if interesting_news:
        random_news = interesting_news[random.randint(0, len(interesting_news)-1)]
    if request.user.is_authenticated:
        friends_count = UserFollowing.objects.filter(user=request.user).count()
    correct_answers = []
    full = []
    for question in history_test_questions:
        answers = [question.answer1, question.answer2, question.answer3, question.correct_answer]
        random.shuffle(answers)
        for n in range(len(answers)):
            if answers[n] == question.correct_answer:
                correct_answers.append(n)
        full.append([question, answers])
    correct_answers = ''.join(list(map(lambda x: str(x), correct_answers)))
    print(correct_answers)
    correct_answers = correct_answers.replace('0', 'A')
    correct_answers = correct_answers.replace('1', 'B')
    correct_answers = correct_answers.replace('2', 'C')
    correct_answers = correct_answers.replace('3', 'D')


    return render(request, "Fsite/problems.html", {'page_obj': page_obj,
                                                   'title': 'Задачи',
                                                   'my_filter': my_filter,
                                                   'friends_count': friends_count,
                                                   'random_news': random_news,
                                                   'materials': materials,
                                                   'subjects': subjects,
                                                   'history': history,
                                                   'history_questions': history_questions,
                                                   'history_test_questions': full,
                                                   'correct_answers': correct_answers})


def problem(request, problem_id):
    post = Problem.objects.filter(pk=problem_id)
    active_contest = Contest.objects.filter(problems__in = post, finished = False).all()
    post = post.first()
    friends_count = 0
    interesting_news = News.objects.filter(is_published = True, tags__slug__in=['interesting_news'])
    random_news = None
    if interesting_news:
        random_news = interesting_news[random.randint(0, len(interesting_news)-1)]
    if (request.user.is_authenticated):
        friends_count = UserFollowing.objects.filter(user=request.user).count()
    if (active_contest):
        post.show_solution = False
    else:
        post.show_solution = True
    post.save()
    return render(request, "Fsite/problem.html", {'problem': post,
                                                  'title': post.title,
                                                  'friends_count': friends_count,
                                                  'random_news': random_news})


def about(request):
    return render(request, "Fsite/about.html", {'title': "О сайте"})


def news_with_id(request, news_id):
    news = News.objects.get(id=news_id)
    friends_count = 0
    interesting_news = News.objects.filter(is_published = True, tags__slug__in=['interesting_news'])
    random_news = None
    if interesting_news:
        random_news = interesting_news[random.randint(0, len(interesting_news)-1)]
    if (request.user.is_authenticated):
        friends_count = UserFollowing.objects.filter(user=request.user).count() 
    return render(request, "Fsite/news_with_id.html", {'news': news,
                                                       'title': news.title,
                                                       'friends_count': friends_count,
                                                       'random_news': random_news})



def news(request):
    is_ajax = request.headers.get('X-Requested-With') == 'XMLHttpRequest'
    paginated_news = Paginator( News.objects.all().order_by('-time_update'), 8) ## тестовый
    if is_ajax:
        response = list(paginated_news.page(request.POST.get('number', 1)).object_list.values())
        return JsonResponse({'response': response, 'page_number': request.POST.get('number', 1)}, safe=False, json_dumps_params={'ensure_ascii': False})
    else:
        return render(request, "Fsite/news_main.html", {'title': 'Новости',
                                                        'news': paginated_news.page(1),
                                                        'paginated_count': paginated_news.num_pages,
                                                        'elements_number': len(paginated_news.page(1).object_list),
                                                        'page_number': request.POST.get('number', 1)})

def problem_in_contest(request, contest_id, problem_id):
    friends_count = 0
    interesting_news = News.objects.filter(is_published = True, tags__slug__in=['interesting_news'])
    random_news = None
    if interesting_news:
        random_news = interesting_news[random.randint(0, len(interesting_news)-1)]
    if (request.user.is_authenticated):
        friends_count = UserFollowing.objects.filter(user=request.user).count() 
    contest = Contest.objects.get(pk=contest_id)
    problem = Problem.objects.get(pk=problem_id)
    if not (problem in contest.problems.all()) or not (request.user in contest.registered_users.all()):
        return redirect('contest', contest_id = contest.id)
    try:
        user_solution = UserSolution.objects.get(user = request.user, contest = contest, problem = problem)
    except ObjectDoesNotExist:
        user_solution = UserSolution(user = request.user, contest = contest, problem = problem)
        user_solution.save()
    if (request.method == "POST"):
        file_path = None
        if user_solution.image_solution:
            file_path = user_solution.image_solution.path
        form = UserSolutionForm(request.POST, request.FILES, instance=user_solution)
        if form.is_valid():
            try:
                form.save()
                if file_path and 'image_solution' in form.changed_data:
                    os.remove(file_path)
                if form['text_solution'].value() == problem.answer:
                    print('ответ правильный')
                    user_solution.mark = UserSolution._meta.get_field('mark').validators[1].limit_value
                    user_solution.save()
                return redirect('contest', contest_id = contest.id)
            except Exception as e:
                print('Ошибка формы ' + str(e))
                form.add_error(e, 'Ошибка добавления формы')
    else:
        form = UserSolutionForm(instance=user_solution)
    return render(request, 'Fsite/problem.html', {'form': form,
                                                  'contest': contest,
                                                  'problem': problem,
                                                  'title': problem.title,
                                                  'friends_count': friends_count,
                                                  'random_news': random_news,
                                                  'user_solution': user_solution})
                                                  


def pageNotFound(request, exception):
    return render(request, "Fsite/error.html", {})


def rating(request):
    users = UserProfile.objects.all().order_by('-rate')
    my_filter = RatingFilter(request.GET, queryset=users)
    users = my_filter.qs
    friends_count = 0
    if (request.user.is_authenticated):
        friends_count = UserFollowing.objects.filter(user=request.user).count() 
    interesting_news = News.objects.filter(is_published = True, tags__slug__in=['interesting_news'])
    random_news = None
    if interesting_news:
        random_news = interesting_news[random.randint(0, len(interesting_news)-1)]
    paginator = Paginator(users, 20) # тестово
    page_number = request.GET.get('page', 1)
    try:
        page_obj = paginator.get_page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.get_page(1)
    except EmptyPage:
        page_obj = paginator.get_page(paginator.num_pages)
    return render(request, "Fsite/rating.html", {'title': "Рейтинг",
                                                 'page_obj': page_obj,
                                                 'friends_count': friends_count, 
                                                 'random_news': random_news,
                                                 "my_filter": my_filter})


class RegisterUser(CreateView):
    form_class = RegisterUserForm
    template_name = 'Fsite/register.html'
    success_url = reverse_lazy('home')
    extra_context = {
        'title': 'Регистрация'
    }

    def form_valid(self, form):
        user = form.save()
        login(self.request, user)
        return redirect('home')


class LoginUser(LoginView):
    form_class = LoginUserForm
    template_name = 'Fsite/login.html'
    extra_context = {
        'title': 'Авторизация'
    }

    def get_success_url(self):
        return reverse_lazy('home')


def change_form(request, username):
    user = UserProfile.objects.filter(username=username).first()
    file_path = user.avatar.path
    file_url = user.avatar.url
    if request.method == "POST":
        if (request.POST.get('user_for_follow', False)):
            user_follow = UserProfile.objects.get(username=request.POST['user_for_follow'])
            current_user = UserProfile.objects.get(username=request.user.username)
            follow = UserFollowing(user=current_user, following_user=user_follow)
            follow.save()
        if (request.POST.get('user_for_unfollow', False)):
            followed_user = UserProfile.objects.get(username=request.POST['user_for_unfollow'])
            current_user = UserProfile.objects.get(username=request.user.username)
            follow = UserFollowing.objects.filter(user=current_user, following_user=followed_user).all()
            follow.delete()
        form = UserProfileChangeForm(request.POST, request.FILES, instance=request.user)
        if form.is_valid():
            try:
                form.save()  
                if file_url != '/media/logo.png' and 'avatar' in form.changed_data:
                    os.remove(file_path)
                return redirect('home')
            except Exception as e:
                form.add_error(None, 'Ошибка добавления формы')
                print('Ошибка формы ' + str(e))
    else:
        form = UserProfileChangeForm(instance=user)
    friends = UserFollowing.objects.filter(user=user).all()
    return render(request, 'Fsite/user_page.html', {"title": username, 
                                                    "form": form,
                                                    "user": user,
                                                    "user_followings": friends})


def contests(request):
    contests = Contest.objects.filter(is_published = True).order_by('completed_at')
    contest_data = {}
    active_contests = Contest.objects.filter(finished = False).order_by('started_at')
    for contest in active_contests:
        started_at = datetime.strftime(contest.started_at, "%Y-%m-%d %H:%M")
        completed_at = datetime.strftime(contest.completed_at, "%Y-%m-%d %H:%M")
        contest_data[contest.pk] = {'started_at': started_at, 'completed_at': completed_at}
    friends_count = 0
    interesting_news = News.objects.filter(is_published = True, tags__slug__in=['interesting_news'])
    random_news = None
    if interesting_news:
        random_news = interesting_news[random.randint(0, len(interesting_news)-1)]
    if (request.user.is_authenticated):
        friends_count = UserFollowing.objects.filter(user=request.user).count() 
    return render(request, 'Fsite/contests.html', {'title': 'Соревнования',
                                                   'contests': contests,
                                                   'active_contests': active_contests,
                                                   'contest_data': json.dumps(contest_data),
                                                   'friends_count': friends_count,
                                                   'random_news': random_news})
    

def contest(request, contest_id):
    contest = Contest.objects.get(pk=contest_id)
    results = UserSolution.objects.filter(contest = contest)
    registered_users = contest.registered_users.all().annotate(mark_sum = Sum("solutions__mark", filter=Q(solutions__contest = contest)))
    registered_users = registered_users.order_by('-mark_sum')
    friends_count = 0
    interesting_news = News.objects.filter(is_published = True, tags__slug__in=['interesting_news'])
    random_news = None
    if interesting_news:
        random_news = interesting_news[random.randint(0, len(interesting_news)-1)]
    if (request.user.is_authenticated):
        friends_count = UserFollowing.objects.filter(user=request.user).count()
    if not (request.user in contest.registered_users.all()) and not contest.is_finished():
        return render(request, 'Fsite/error_403.html', {}) 
    return render(request, 'Fsite/contest.html', {'title': 'Соревнование',
                                                  'contest': contest,
                                                  'results': results,
                                                  'registered_users': registered_users,
                                                  'random_news': random_news,
                                                  'friends_count': friends_count})

def logout_user(request):
    logout(request)
    return redirect('home')

def registration_for_contest(request, contest_id, username):
    try:
        user = UserProfile.objects.get(username=username)
    except ObjectDoesNotExist:
        return redirect('login')
    contest = Contest.objects.get(pk=contest_id)

    if ( contest.is_running() or contest.is_finished() ):
        return redirect('contest', contest_id)
    if (user == request.user):
        contest.registered_users.add(user)
        contest.save()
    return redirect('contests')


def checking_contest_solutions(request, contest_id):
    friends_count = 0
    contest = Contest.objects.get(pk=contest_id)
    interesting_news = News.objects.filter(is_published = True, tags__slug__in=['interesting_news'])
    interesting_news = News.objects.filter(is_published = True, tags__slug__in=['interesting_news'])
    random_news = None
    queryset = UserSolution.objects.filter(contest=contest)
    SolutionVerificationFormSet = modelformset_factory(UserSolution, form = SolutionVerification,fields=('mark',), extra=0)
    forms = SolutionVerificationFormSet(queryset=queryset)
    if request.method == "POST":
        forms = SolutionVerificationFormSet(request.POST)
        if forms.is_valid():
            try:
                forms.save()
                print("сохранение формы прошло успешно")
            except Exception as e:
                print(e)
                forms.add_error(e)
    if interesting_news:
        random_news = interesting_news[random.randint(0, len(interesting_news)-1)]
    if (request.user.is_authenticated):
        friends_count = UserFollowing.objects.filter(user=request.user).count() 
    if (request.user in contest.judges.all()):
        return render(request, 'Fsite/checking_contest_solutions.html', {'title': 'Проверка контеста',
                                                                        'contest': contest,
                                                                        'friends_count': friends_count,
                                                                        'random_news': random_news,
                                                                        'forms': forms
                                                                        })
    else:
        return render(request, 'Fsite/error_403.html', {})


def render_find_info(request):
    query_dict = request.GET
    news = News.objects.none()
    contests = Contest.objects.none()
    problems = Problem.objects.none()
    users = UserProfile.objects.none()
    if ('q' in query_dict):
        query = query_dict['q'].replace(' ', '')
        news = News.objects.filter(Q(title__iregex = query) | Q(content__iregex = query)).order_by('-time_update')
        contests = Contest.objects.filter(name__iregex = query).order_by('-completed_at')
        problems = Problem.objects.filter(title__iregex = query).order_by('rate')
        users = UserProfile.objects.filter(username__iregex = query)
    return render (request, 'Fsite/find_information.html', {'news': news,
                                                            'contests': contests,
                                                            'problems': problems,
                                                            'users': users, 
                                                            'title': 'Поиск'})


def download_history_tests(request):
    document = Document()

    history_test_questions_num = request.GET.get('history_test_questions_number', None)
    history_test_cases = request.GET.get('history_test_cases', None)
    history_test_questions = HistoryTestProblem.objects.none()
    history_test_questions_all = HistoryTestProblem.objects.filter(is_published=True).all()
    if history_test_cases:
        if history_test_questions_num:
            history_test_questions_num = int(history_test_questions_num)
            for i in range(int(history_test_cases)):
                document.add_heading(f'Вариант {i + 1}', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

                if history_test_questions_num < len(history_test_questions_all):
                    history_test_questions = random_questions(history_test_questions_num, history_test_questions_all, HistoryTestProblem)
                else:
                    history_test_questions = history_test_questions_all.all()
                correct_answers = []
                for question in history_test_questions:
                    answers = [question.answer1, question.answer2, question.answer3, question.correct_answer]
                    random.shuffle(answers)
                    for n in range(len(answers)):
                        if answers[n] == question.correct_answer:
                            correct_answers.append(n)
                    document.add_heading(question.question, 1)
                    document.add_paragraph(f'A) {answers[0]}')
                    document.add_paragraph(f'B) {answers[1]}')
                    document.add_paragraph(f'C) {answers[2]}')
                    document.add_paragraph(f'D) {answers[3]}')
                correct_answers = ''.join(list(map(lambda x: str(x), correct_answers)))
                correct_answers = correct_answers.replace('0', 'A')
                correct_answers = correct_answers.replace('1', 'B')
                correct_answers = correct_answers.replace('2', 'C')
                correct_answers = correct_answers.replace('3', 'D')
                paragraph = document.add_paragraph()
                run = paragraph.add_run('Ответы: ')
                run.font.color.rgb = RGBColor(255, 0, 0)
                run = paragraph.add_run(f"{correct_answers}")
                run.font.color.rgb = RGBColor(0, 0, 0)


    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=download.docx'
    document.save(response)
    return response


# class UserInfo(LoginView):
#     form_class = UserProfileChangeForm
#     template_name = 'Fsite/login.html'
#     user = UserProfile.objects.filter(username=username).first()
#     context = {
#         'title': 'Авторизация'
#         ''
#     }

#     def get_success_url(self):
#         return reverse_lazy('home')
